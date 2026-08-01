from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("artifacts/2026-7-31-例会作业.docx")

BLACK = RGBColor(0, 0, 0)
BLUE = BLACK
DARK_BLUE = BLACK
NAVY = BLACK
GRAY = BLACK

ASCII_FONT = "Arial Unicode MS"
CJK_FONT = "Arial Unicode MS"


def set_run_font(
    run,
    *,
    size=11,
    color=BLACK,
    bold=False,
    italic=False,
):
    run.font.name = ASCII_FONT
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_style_font(style, size, color=BLACK, bold=False):
    style.font.name = ASCII_FONT
    style._element.get_or_add_rPr()
    fonts = style._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, 14, BLACK, True)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.line_spacing = 1.0
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, 12, BLACK, True)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(5)
    heading2.paragraph_format.line_spacing = 1.0
    heading2.paragraph_format.keep_with_next = True

    heading3 = doc.styles["Heading 3"]
    set_style_font(heading3, 11, BLACK, True)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)
    heading3.paragraph_format.line_spacing = 1.0
    heading3.paragraph_format.keep_with_next = True


def add_numbering_definition(doc, *, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)

    if bullet:
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Symbol")
        rfonts.set(qn("w:hAnsi"), "Symbol")
        rpr.append(rfonts)
        level.append(rpr)

    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_numbered_item(doc, num_id, label, body):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.paragraph_format.keep_together = True
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run)
    return paragraph


def add_bullet(doc, num_id, text):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    set_run_font(field_run, size=9, color=GRAY)
    field_run._r.append(begin)
    field_run._r.append(instr)
    field_run._r.append(separate)
    field_run._r.append(value)
    field_run._r.append(end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=GRAY)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    number_id = add_numbering_definition(doc, bullet=False)
    bullet_id = add_numbering_definition(doc, bullet=True)

    # Plain black-and-white opening that resembles a default Word assignment.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("例会作业")
    set_run_font(title_run, size=16, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_run = subtitle.add_run("个性化多模式路线推荐——用户画像部分")
    set_run_font(subtitle_run, size=11, color=BLACK)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(0)
    date.paragraph_format.space_after = Pt(12)
    date_run = date.add_run("2026年7月31日")
    set_run_font(date_run, size=10, color=BLACK)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    intro_run = intro.add_run(
        "这两天主要在看个性化多模式路线推荐中的用户画像问题，重点想弄清楚画像里要放哪些信息、偏好权重怎样得到，以及路线的时间、费用等属性从哪里来。目前完成的工作如下。"
    )
    set_run_font(intro_run)

    doc.add_heading("一、已完成的工作", level=1)

    add_numbered_item(
        doc,
        number_id,
        "查找并整理相关论文。 ",
        "围绕用户画像、情境感知、动态偏好和个性化多模式路线推荐整理了一批论文，又补充查找了10篇质量较高的相关论文，后面会继续筛选和阅读。",
    )

    paragraph = add_numbered_item(
        doc,
        number_id,
        "重点阅读画像权重相关方法。 ",
        "重点看了《Preference-Aware Multimodal Journey Planner》和FAVOUR。前一篇不是让用户直接填“时间权重是多少”，而是让用户评价几组特点不同的路线，再反推出时间、费用、步行、换乘等权重；FAVOUR则采用“群体偏好作初始值—路线比较学习个人偏好—根据真实选择继续更新”的流程。",
    )

    add_numbered_item(
        doc,
        number_id,
        "重新梳理25号公铁联程论文。 ",
        "这篇论文通过历史订单计算旅客在出发时段、列车类型等方面的偏好，再把旅客分成价格敏感、时间敏感等类型。它的交通方式比较少，不能直接套用到多模式接驳，但“历史行为偏好＋指标敏感性”的思路可以参考。",
    )

    add_numbered_item(
        doc,
        number_id,
        "初步确定第一版画像的结构。 ",
        "目前把画像分成长期基础画像、本次出行状态和实时环境三部分。长期部分记录稳定偏好和限制；本次出行记录是否赶时间、预算、行李等；实时环境记录天气、延误、拥堵和交通方式是否可用。这样不会因为一次下雨或一次赶时间就直接改变用户的长期画像。",
    )

    add_numbered_item(
        doc,
        number_id,
        "确定第一版的路线评价指标。 ",
        "先使用时间、费用、步行距离和换乘次数四个指标。新用户先通过少量路线比较得到基础权重，本次出行和实时环境只对这些权重做临时调整。不能接受的交通方式、最大步行距离、最大换乘次数等先作为硬限制筛选路线，不只靠降低分数处理。",
    )

    add_numbered_item(
        doc,
        number_id,
        "梳理路线属性的获取方式。 ",
        "路线规划服务不能只返回一条路线名称或地图线条，还要给出每一段的交通方式、距离、乘车时间、等待时间和换乘信息。系统再汇总出总时间、总费用、总步行距离和换乘次数；票价、延误和车辆可用性不足时，需要再接入票价规则或实时交通数据。",
    )

    add_numbered_item(
        doc,
        number_id,
        "整理阅读记录。 ",
        "已经把关注用户画像的论文按“核心思路—具体流程—可以参考的地方—局限”整理到Markdown文件中，同时把公式改成了可以直接渲染的格式，方便后面继续补充。",
    )

    doc.add_heading("二、目前得到的结论", level=1)

    add_bullet(
        doc,
        bullet_id,
        "“先生成多样化候选路线，再根据用户偏好排序”这个方向是可行的，但前面的候选路线必须有明显差异，不能全是时间和费用都差不多的路线。",
    )
    add_bullet(
        doc,
        bullet_id,
        "用户画像不能只用年龄、职业等静态标签，更重要的是用户对时间、费用、步行和换乘的偏好，以及本次出行和实时环境带来的临时变化。",
    )
    add_bullet(
        doc,
        bullet_id,
        "第一版先采用简单、能解释的方法，不急着加入Q-learning、BERT或联邦学习。先验证权重能不能学出来、路线排序是否合理，再考虑复杂模型。",
    )

    doc.add_heading("三、下一步要做的事情", level=1)

    add_bullet(
        doc,
        bullet_id,
        "设计6～8组路线比较题，覆盖时间与费用、时间与步行、时间与换乘等取舍关系。",
    )
    add_bullet(
        doc,
        bullet_id,
        "确定反推四个基础权重的具体计算方法，并用小规模示例验证结果是否符合直觉。",
    )
    add_bullet(
        doc,
        bullet_id,
        "了解现有路线规划服务能返回哪些字段，缺少的费用和实时信息再确定数据来源。",
    )
    add_bullet(
        doc,
        bullet_id,
        "为赶时间、预算有限、携带行李、下雨等情况先设置一组透明的调整规则，后面再根据真实数据修改。",
    )

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    add_page_number(footer)

    doc.core_properties.title = "例会作业"
    doc.core_properties.subject = "个性化多模式路线推荐中的用户画像阶段工作记录"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
